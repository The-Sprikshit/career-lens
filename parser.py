"""
Resume Parser — extracts structured info from PDF/DOCX/TXT resumes.

Robust file format handling:
- Accepts file paths OR file-like objects (BytesIO, UploadedFile)
- Tries PDF → DOCX → TXT with graceful fallback
- Returns clear error messages if a format fails
- No hardcoded /tmp/ paths (cross-platform)
"""

import re
import io
from pathlib import Path
from typing import Dict, Union, BinaryIO


# ── Lazy imports (don't crash if libs missing) ──────────────────────
def _safe_import_pdfplumber():
    try:
        import pdfplumber
        return pdfplumber
    except ImportError:
        return None

def _safe_import_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        return None


# ── Text extraction ─────────────────────────────────────────────────
def extract_text_from_pdf(source: Union[str, BinaryIO]) -> str:
    """Extract text from PDF (file path OR file-like object)."""
    pdfplumber = _safe_import_pdfplumber()
    if pdfplumber is None:
        raise ImportError(
            "pdfplumber not installed. Run: pip install pdfplumber"
        )

    text_parts = []
    try:
        if isinstance(source, (str, Path)):
            ctx = pdfplumber.open(str(source))
        else:
            # file-like object (BytesIO, UploadedFile)
            ctx = pdfplumber.open(source)

        with ctx as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                text_parts.append(txt)
    except Exception as e:
        raise RuntimeError(f"PDF parsing failed: {e}") from e

    return "\n".join(text_parts)


def extract_text_from_docx(source: Union[str, BinaryIO]) -> str:
    """Extract text from DOCX (file path OR file-like object)."""
    Document = _safe_import_docx()
    if Document is None:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )

    try:
        if isinstance(source, (str, Path)):
            doc = Document(str(source))
        else:
            doc = Document(source)

        return "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:
        raise RuntimeError(f"DOCX parsing failed: {e}") from e


def extract_text(source: Union[str, Path, BinaryIO],
                 filename_hint: str = None) -> str:
    """
    Smart dispatcher — tries formats based on file extension OR content.
    Always falls back gracefully.

    Args:
        source: file path OR file-like object
        filename_hint: filename for extension detection (e.g. "resume.pdf")
    """
    # Determine format
    if filename_hint:
        ext = Path(filename_hint).suffix.lower()
    elif isinstance(source, (str, Path)):
        ext = Path(source).suffix.lower()
    else:
        ext = ""

    # Reset file pointer if file-like
    if hasattr(source, "seek"):
        try:
            source.seek(0)
        except Exception:
            pass

    # Dispatch
    if ext == ".pdf":
        return extract_text_from_pdf(source)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(source)
    elif ext in (".txt", ".md", ""):
        if isinstance(source, (str, Path)):
            return Path(source).read_text(encoding="utf-8", errors="ignore")
        else:
            content = source.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="ignore")
            return content
    else:
        # Try PDF first, then DOCX, then text
        for extractor, label in [
            (extract_text_from_pdf, "PDF"),
            (extract_text_from_docx, "DOCX"),
        ]:
            try:
                if hasattr(source, "seek"):
                    source.seek(0)
                return extractor(source)
            except Exception:
                continue
        # Last resort: treat as text
        if isinstance(source, (str, Path)):
            return Path(source).read_text(encoding="utf-8", errors="ignore")
        content = source.read()
        return content.decode("utf-8", errors="ignore") if isinstance(content, bytes) else content


# ── Regex patterns ──────────────────────────────────────────────────
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\s\-\(\)]{8,}\d)")
URL_RE = re.compile(r"(https?://[^\s]+|www\.[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+)")
DATE_RE = re.compile(
    r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]+\d{4}\b"
    r"|\b\d{1,2}/\d{4}\b"
    r"|\b\d{4}\s*[–\-to]+\s*(?:\d{4}|Present|Current|Now)\b",
    re.IGNORECASE,
)
SECTION_HEADERS = [
    "education", "experience", "work experience", "professional experience",
    "skills", "technical skills", "projects", "certifications",
    "achievements", "summary", "objective", "profile", "contact",
    "publications", "awards", "languages", "interests", "hobbies",
    "volunteer", "extracurricular",
]
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+", re.IGNORECASE)
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+", re.IGNORECASE)

# Words that show up on the first non-blank line but are never a person's name
_NAME_BLOCKLIST = {
    "resume", "curriculum vitae", "cv", "portfolio", "profile",
    *SECTION_HEADERS,
}


def _extract_name(text: str, email: str = None) -> str:
    """
    Best-effort candidate name extraction.

    Strategy (first match wins):
      1. First non-empty line at the top of the file, if it looks like a
         name (2-4 title-case words, no digits/@/urls, not a section header).
      2. Fall back to a cleaned-up version of the email's local part.
      3. Give up and return None — callers should handle this gracefully.
    """
    for raw_line in text.strip().split("\n")[:8]:
        line = raw_line.strip().strip("|•-").strip()
        if not line:
            continue
        line_lower = line.lower()
        if line_lower in _NAME_BLOCKLIST:
            continue
        if "@" in line or re.search(r"https?://|www\.", line):
            continue
        if any(ch.isdigit() for ch in line):
            continue
        words = line.split()
        if not (1 < len(words) <= 4):
            continue
        if len(line) > 40:
            continue
        # Looks like "John Doe" / "JOHN DOE" / "John A. Doe"
        if all(w[0].isalpha() and (w[0].isupper() or w.isupper()) for w in words):
            return " ".join(w.capitalize() if w.isupper() and len(w) > 1 else w
                            for w in words)

    if email:
        local = email.split("@")[0]
        local = re.sub(r"[._\-\d]+", " ", local).strip()
        if local:
            return " ".join(part.capitalize() for part in local.split())

    return None


def parse_resume(source: Union[str, Path, BinaryIO],
                 filename_hint: str = None) -> Dict:
    """
    Parse a resume and return structured info.

    Accepts:
      - file path (str / Path)
      - file-like object (BytesIO, Streamlit UploadedFile, etc.)

    Returns:
    {
        "raw_text": "...", "email": "...", "phone": "...",
        "links": [...], "sections": {...},
        "years_experience": 3.5, "education_count": 2,
        "job_titles": [...], "word_count": 450,
        "parse_warnings": [...]   # any format issues
    }
    """
    warnings = []

    try:
        text = extract_text(source, filename_hint=filename_hint)
    except Exception as e:
        warnings.append(f"Text extraction failed: {e}")
        text = ""

    if not text.strip():
        warnings.append("No text extracted from file")

    text_lower = text.lower()

    # Basic contact info
    email = None
    m = EMAIL_RE.search(text)
    if m:
        email = m.group(0)

    phone = None
    m = PHONE_RE.search(text)
    if m:
        phone = m.group(1).strip()

    links = list(set(URL_RE.findall(text)))

    m = GITHUB_RE.search(text)
    github = m.group(0) if m else None
    m = LINKEDIN_RE.search(text)
    linkedin = m.group(0) if m else None

    name = _extract_name(text, email)

    # Sections
    sections = _split_into_sections(text)

    # Years of experience from date ranges
    dates = DATE_RE.findall(text)
    years_exp = _estimate_years_experience(dates, text)

    # Education count
    edu_keywords = ["bachelor", "master", "phd", "b.tech", "m.tech", "mba",
                    "degree", "university", "college"]
    edu_count = sum(1 for kw in edu_keywords if kw in text_lower)
    edu_count = max(1, edu_count // 2)

    # Job titles
    from data.role_profiles import get_role_list_flat
    common_titles = get_role_list_flat()
    job_titles = [t for t in common_titles if t in text_lower]
    job_titles = _dedupe_titles(job_titles)

    return {
        "raw_text": text,
        "name": name,
        "email": email,
        "phone": phone,
        "github": github,
        "linkedin": linkedin,
        "links": links,
        "sections": sections,
        "years_experience": years_exp,
        "education_count": edu_count,
        "job_titles": job_titles,
        "word_count": len(text.split()),
        "parse_warnings": warnings,
    }


def _split_into_sections(text: str) -> Dict[str, str]:
    """Split resume text into sections by headers."""
    sections: Dict[str, str] = {}
    lines = text.split("\n")
    current_section = "header"
    buffer = []

    for line in lines:
        line_stripped = line.strip()
        line_lower = line_stripped.lower()
        is_header = (
            line_stripped in [h.title() for h in SECTION_HEADERS]
            or line_lower in SECTION_HEADERS
            or (len(line_stripped) < 40 and line_lower in SECTION_HEADERS)
        )
        if is_header:
            if buffer:
                sections[current_section] = "\n".join(buffer).strip()
            current_section = line_lower
            buffer = []
        else:
            buffer.append(line)
    if buffer:
        sections[current_section] = "\n".join(buffer).strip()

    return sections


def _estimate_years_experience(dates, text: str) -> float:
    """Estimate total years of experience from date ranges."""
    import datetime
    total_years = 0.0
    now = datetime.datetime.now().year
    for d in dates:
        years = re.findall(r"\b(?:19|20)\d{2}\b", d)
        if len(years) >= 2:
            try:
                start = int(years[0]); end = int(years[1])
                if end >= start and (now - end) < 50:
                    total_years += (end - start)
            except Exception:
                pass
    if total_years == 0:
        all_years = [int(y) for y in re.findall(r"\b(?:19|20)\d{2}\b", text)]
        if all_years:
            total_years = max(0, max(all_years) - min(all_years))
    return round(min(total_years, 40), 1)


def _dedupe_titles(titles: list) -> list:
    """Remove titles that are substrings of others."""
    result = []
    titles_sorted = sorted(titles, key=len, reverse=True)
    for t in titles_sorted:
        if not any(t in r and t != r for r in result):
            result.append(t)
    return result


# ── CLI test ────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys, json
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        display = {k: v for k, v in result.items() if k != "raw_text"}
        print(json.dumps(display, indent=2, default=str))
    else:
        # Demo: parse from stdin / inline string
        sample = """
        John Doe
        john@example.com | linkedin.com/in/johndoe

        EXPERIENCE
        Software Engineer (2020 - 2023)
        • Built Python APIs
        """
        result = parse_resume(io.StringIO(sample), filename_hint="test.txt")
        display = {k: v for k, v in result.items() if k != "raw_text"}
        print(json.dumps(display, indent=2, default=str))
